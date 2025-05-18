import streamlit as st
import faiss
import os
from io import BytesIO
from docx import Document
import numpy as np
from langchain_community.document_loaders import WebBaseLoader
from PyPDF2 import PdfReader
from langchain.chains import RetrievalQA
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_huggingface import HuggingFaceEndpoint

# Load the API key securely from Streamlit secrets
huggingface_api_key = st.secrets["HUGGINGFACEHUB_API_TOKEN"]
os.environ['HUGGINGFACEHUB_API_TOKEN'] = huggingface_api_key

def process_input(input_type, input_data):
    loader = None
    if input_type == "Link":
        loader = WebBaseLoader(input_data)
        documents = loader.load()
    elif input_type == "PDF":
        if isinstance(input_data, BytesIO):
            pdf_reader = PdfReader(input_data)
        else:
            pdf_reader = PdfReader(BytesIO(input_data.read()))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        documents = text
    elif input_type == "Text":
        documents = input_data
    elif input_type == "DOCX":
        if isinstance(input_data, BytesIO):
            doc = Document(input_data)
        else:
            doc = Document(BytesIO(input_data.read()))
        text = "\n".join([para.text for para in doc.paragraphs])
        documents = text
    elif input_type == "TXT":
        if isinstance(input_data, BytesIO):
            text = input_data.read().decode('utf-8')
        else:
            text = input_data.read().decode('utf-8')
        documents = text
    else:
        raise ValueError("Unsupported input type")

    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    if input_type == "Link":
        texts = text_splitter.split_documents(documents)
        texts = [str(doc.page_content) for doc in texts]
    else:
        texts = text_splitter.split_text(documents)

    hf_embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-mpnet-base-v2",
        model_kwargs={'device': 'cpu'},
        encode_kwargs={'normalize_embeddings': False}
    )

    sample_embedding = np.array(hf_embeddings.embed_query("sample text"))
    dimension = sample_embedding.shape[0]
    index = faiss.IndexFlatL2(dimension)

    vector_store = FAISS(
        embedding_function=hf_embeddings.embed_query,
        index=index,
        docstore=InMemoryDocstore(),
        index_to_docstore_id={},
    )
    vector_store.add_texts(texts)
    return vector_store

def answer_question(vectorstore, query):
    llm = HuggingFaceEndpoint(
        repo_id='meta-llama/Meta-Llama-3-8B-Instruct',
        token=huggingface_api_key,
        temperature=0.6
    )
    qa = RetrievalQA.from_chain_type(llm=llm, retriever=vectorstore.as_retriever())
    return qa({"query": query})

def main():
    st.title("RAG Q&A App")
    input_type = st.selectbox("Input Type", ["Link", "PDF", "Text", "DOCX", "TXT"])
    if input_type == "Link":
        number_input = st.number_input(min_value=1, max_value=20, step=1, label="Enter the number of Links")
        input_data = []
        for i in range(number_input):
            url = st.sidebar.text_input(f"URL {i+1}")
            input_data.append(url)
    elif input_type == "Text":
        input_data = st.text_input("Enter the text")
    elif input_type == 'PDF':
        input_data = st.file_uploader("Upload a PDF file", type=["pdf"])
    elif input_type == 'TXT':
        input_data = st.file_uploader("Upload a text file", type=['txt'])
    elif input_type == 'DOCX':
        input_data = st.file_uploader("Upload a DOCX file", type=['docx', 'doc'])

    if st.button("Proceed"):
        vectorstore = process_input(input_type, input_data)
        st.session_state["vectorstore"] = vectorstore

    if "vectorstore" in st.session_state:
        query = st.text_input("Ask your question")
        if st.button("Submit"):
            answer = answer_question(st.session_state["vectorstore"], query)
            st.write(answer)

if __name__ == "__main__":
    main()
